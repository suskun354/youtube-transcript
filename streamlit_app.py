import streamlit as st
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api._errors import TranscriptsDisabled, NoTranscriptFound, VideoUnavailable, TooManyRequests
import docx
import io
import base64
import re
import requests
import time
from pytube import YouTube
import os

def extract_video_id(url):
    """YouTube URL'sinden video ID'sini çıkarır."""
    # Normal URL formatı (youtube.com/watch?v=VIDEO_ID)
    video_id_match = re.search(r'(?:v=|\/)([0-9A-Za-z_-]{11}).*', url)
    if video_id_match:
        return video_id_match.group(1)
    
    # Kısa URL formatı (youtu.be/VIDEO_ID)
    video_id_match = re.search(r'youtu\.be\/([0-9A-Za-z_-]{11})', url)
    if video_id_match:
        return video_id_match.group(1)
    
    return None

def get_transcript_pytube(video_url):
    """Pytube kütüphanesi kullanarak altyazı alır."""
    try:
        yt = YouTube(video_url)
        caption_tracks = yt.captions
        
        # Önce Türkçe veya İngilizce altyazıları deneyelim
        caption = None
        for c in caption_tracks:
            code = c.code.split('.')[0]  # 'a.en' -> 'a'
            if code in ['tr', 'en']:
                caption = c
                break
        
        # Eğer bulamazsak, ilk altyazıyı kullanalım
        if caption is None and len(caption_tracks) > 0:
            caption = caption_tracks[0]
        
        if caption is None:
            return "Hata: Bu video için altyazı bulunamadı.", None
        
        # XML altyazıyı alalım ve metni çıkaralım
        xml_captions = caption.xml_captions
        
        # Basit bir XML parser ile metni çıkaralım
        # Bu çok basit bir parser - gerçek kullanımda daha gelişmiş bir XML parser kullanılmalı
        clean_text = re.sub(r'<.*?>', '', xml_captions)
        clean_text = re.sub(r'\n', ' ', clean_text)
        
        return clean_text, extract_video_id(video_url)
    except Exception as e:
        return f"Pytube ile altyazı alınırken hata: {str(e)}", None

def get_transcript(video_url):
    """Ana transkript alma fonksiyonu - birden fazla yöntem dener."""
    try:
        video_id = extract_video_id(video_url)
        if not video_id:
            return "Hata: Geçerli bir YouTube URL'si değil", None
        
        # Rate limiting sorunlarını azaltmak için biraz bekleme ekleyelim
        time.sleep(1)
        
        # Önce İngilizce ve Türkçe dil seçenekleriyle deneyelim
        try:
            transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=['en', 'tr'])
            raw_text = " ".join([entry['text'] for entry in transcript])
            return raw_text, video_id
        except Exception as youtube_api_error:
            # YouTube API başarısız olursa pytube'u deneyelim
            st.warning("YouTube Transcript API ile transkript alınamadı. Alternatif yöntem deneniyor...")
            transcript_text, video_id = get_transcript_pytube(video_url)
            
            if not transcript_text.startswith("Hata") and not transcript_text.startswith("Pytube"):
                return transcript_text, video_id
            else:
                # Pytube da başarısız olursa, ilk hatayı detaylandırarak döndürelim
                if isinstance(youtube_api_error, TranscriptsDisabled):
                    return "Hata: Bu video için altyazılar devre dışı bırakılmış.", None
                elif isinstance(youtube_api_error, NoTranscriptFound):
                    return "Hata: Bu video için transkript bulunamadı. Video sahibi tarafından eklenmiş altyazı olmayabilir.", None
                elif isinstance(youtube_api_error, VideoUnavailable):
                    return "Hata: Video kullanılamıyor veya özel olabilir.", None
                elif isinstance(youtube_api_error, TooManyRequests):
                    return "Hata: YouTube API'den çok fazla istek yapıldı. Lütfen daha sonra tekrar deneyin veya IP kısıtlamaları nedeniyle videoyu yerel olarak indirmeyi deneyin.", None
                else:
                    return f"Beklenmeyen bir hata oluştu: {str(youtube_api_error)}", None
    except Exception as e:
        return f"Beklenmeyen bir hata oluştu: {str(e)}", None

def create_word_document(transcript, video_id):
    """Transkripti Word dosyasına dönüştürür."""
    doc = docx.Document()
    doc.add_heading(f'YouTube Video Transkripti - {video_id}', 0)
    
    p = doc.add_paragraph(transcript)
    
    # Belgeyi byte olarak kaydet
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def get_download_link(buffer, filename):
    """Word dosyasını indirmek için bir link oluşturur."""
    b64 = base64.b64encode(buffer.getvalue()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}.docx">Word Dosyası Olarak İndir</a>'

# Sayfa yapılandırması ve başlık
st.set_page_config(
    page_title="YouTube Video Transkripti",
    page_icon="🎬",
    layout="wide"
)

# Başlık ve açıklama
st.title("YouTube Video Transkripti")
st.markdown("YouTube video URL'sini girin ve transkripti alın.")

# IP uyarısı
st.info("⚠️ Not: YouTube API sorunları durumunda otomatik olarak alternatif yöntemler denenecektir. Herhangi bir sorun yaşarsanız, uygulamayı yerel bilgisayarınızda çalıştırmayı deneyin.")

# Ana içerik
col1, col2 = st.columns([1, 2])

with col1:
    video_url = st.text_input("Video URL", placeholder="YouTube Video URL'sini buraya yapıştırın...")
    get_btn = st.button("Transkripti Al", type="primary")

with col2:
    # Video gösterimi için bir yer ayırıyoruz
    if video_url:
        try:
            video_id = extract_video_id(video_url)
            if video_id:
                st.video(f"https://youtu.be/{video_id}")
            else:
                st.error("Geçerli bir YouTube URL'si giriniz.")
        except:
            st.error("Geçerli bir YouTube URL'si giriniz.")

# Transkript sonucu için büyük bir alan
transcript_container = st.container()

# Butona tıklandığında transkript al
if get_btn:
    if not video_url:
        st.error("Lütfen bir YouTube video URL'si girin.")
    else:
        with st.spinner("Transkript alınıyor..."):
            transcript, video_id = get_transcript(video_url)
        
        with transcript_container:
            st.text_area("Video Transkripti", value=transcript, height=400)
            
            # Transkript alındıysa ve hata yoksa Word olarak indirme seçeneği göster
            if video_id and not transcript.startswith("Hata:") and not transcript.startswith("Beklenmeyen"):
                word_buffer = create_word_document(transcript, video_id)
                
                # İndirme düğmesi
                st.download_button(
                    label="Word Dosyası Olarak İndir",
                    data=word_buffer,
                    file_name=f"transkript_{video_id}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# IP sorunları için çözüm önerileri
expander = st.expander("YouTube API erişim sorunları yaşıyorsanız tıklayın")
with expander:
    st.markdown("""
    ## Olası Çözümler
    
    YouTube API bazen IP kısıtlamaları uygulayabilir, özellikle bulut ortamlarında çalışırken.
    
    Alternatif çözümler:
    
    1. **Uygulamayı yerel bilgisayarınızda çalıştırın**:
       ```
       pip install -r requirements.txt
       streamlit run app.py
       ```
       
    2. **VPN kullanın**: IP adresinizi değiştirmek sorunu çözebilir.
    
    3. **YouTube Premium aboneliği** ile bazı kısıtlamalar azaltılabilir.
    
    4. **Uygulama artık otomatik olarak alternatif yöntemleri dener**: İlk yöntem başarısız olursa, uygulama otomatik olarak pytube kütüphanesini kullanarak transkript almayı dener.
    """)

# Alt bilgi
st.markdown("---")
st.markdown("YouTube Transkript Uygulaması | [GitHub](https://github.com/suskun354/youtube-transcript)")
